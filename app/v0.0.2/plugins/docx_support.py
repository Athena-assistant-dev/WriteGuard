
from flask import Blueprint, request, jsonify, send_from_directory
import os
import base64
from docx import Document

file_docx = Blueprint("file_docx", __name__)

# Use the same BASE_PATH environment variable as the main API
BASE_DIR = os.environ.get("BASE_PATH", "")

@file_docx.route("/docx/download/<path:filename>", methods=["GET"])
def download_docx(filename):
    return send_from_directory(BASE_DIR, filename, as_attachment=True)

@file_docx.route("/docx/read-base64", methods=["POST"])
def read_docx_base64():
    from smart_safe_write import auto_map_on_failure

    @auto_map_on_failure
    def wrapped():
        filepath = request.json.get("filepath")
        full_path = os.path.join(BASE_DIR, filepath)
        if not os.path.exists(full_path):
            return jsonify({"success": False, "error": "File not found"})
        with open(full_path, "rb") as f:
            encoded = base64.b64encode(f.read()).decode("utf-8")
        return jsonify({"success": True, "base64": encoded})

    return wrapped()

@file_docx.route("/docx/parse", methods=["POST"])
def parse_docx():
    from smart_safe_write import auto_map_on_failure

    @auto_map_on_failure
    def wrapped():
        try:
            data = request.get_json()
            filepath = data.get("filepath")
            full_path = os.path.abspath(os.path.join(BASE_DIR, filepath))

            if not os.path.exists(full_path):
                return jsonify({"error": "File not found", "success": False})

            doc = Document(full_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return jsonify({"success": True, "text": text})

        except Exception as e:
            return jsonify({"error": "Internal server error", "details": str(e)})

    return wrapped()

@file_docx.route("/docx/write", methods=["POST"])
def write_docx():
    from smart_safe_write import smart_safe_write, auto_map_on_failure

    @auto_map_on_failure
    def wrapped():
        try:
            data = request.get_json()
            filepath = data.get("filepath")
            content = data.get("content")

            if not filepath or not content:
                return jsonify({"error": "Missing filepath or content", "success": False}), 400

            full_path = os.path.abspath(os.path.join(BASE_DIR, filepath))

            doc = Document()
            for line in content.splitlines():
                doc.add_paragraph(line)

            tmp_path = full_path + ".tmp"
            doc.save(tmp_path)

            with open(tmp_path, "rb") as f:
                # The content is passed as bytes, so mode='binary' is not needed and incorrect.
                result = smart_safe_write(full_path, content_bytes=f.read())

            os.remove(tmp_path)

            return jsonify({"success": True, "result": result})

        except Exception as e:
            return jsonify({"error": str(e), "success": False}), 500

    return wrapped()
